import braille
import speech_recognition as sr
from docx import Document

def sign2text(text):
    doc = Document()
    for char in text:
        char = char.lower()
        if char == "a":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\A.PNG")
        elif char == "b":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\B.PNG")
        elif char == "c":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\C.PNG")
        elif char == "d":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\D.PNG")
        elif char == "e":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\E.PNG")
        elif char == "f":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\F.PNG")
        elif char == "g":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\G.PNG")
        elif char == "h":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\H.PNG")
        elif char == "i":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\I.PNG")
        elif char == "j":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\J.PNG")
        elif char == "k":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\K.PNG")
        elif char == "l":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\L.PNG")
        elif char == "m":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\M.PNG")
        elif char == "n":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\N.PNG")
        elif char == "o":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\O.PNG")
        elif char == "p":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\P.PNG")
        elif char == "q":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\Q.PNG")
        elif char == "r":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\R.PNG")
        elif char == "s":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\S.PNG")
        elif char == "t":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\T.PNG")
        elif char == "u":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\U.PNG")
        elif char == "v":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\V.PNG")
        elif char == "w":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\W.PNG")
        elif char == "x":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\X.PNG")
        elif char == "y":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\Y.PNG")
        elif char == "z":
            doc.add_picture(r"C:\Users\joyrb\PycharmProjects\inOut_project\letters\Z.PNG")
        elif char == " ":
            doc.add_paragraph(' ')

        doc.save('SignTranscript.docx')


r = sr.Recognizer()
transcript = Document()
brailleScript = Document()

try:

    with sr.Microphone() as source2:

        r.adjust_for_ambient_noise(source2, duration=0.2)
        audio2 = r.listen(source2)

        MyText = r.recognize_google(audio2)
        MyText = MyText.lower()

        print(MyText)
        transcript.add_paragraph(MyText)
        transcript.save('transcript.docx')

        brailleTranslation = braille.textToBraille(MyText)
        brailleScript.add_paragraph(brailleTranslation)
        brailleScript.save('brailleScript.docx')

        sign2text(MyText)


except sr.RequestError as e:
    print("Could not request results; {0}".format(e))

except sr.UnknownValueError:
    print("unknown error occured")